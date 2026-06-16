from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "outputs"
ASSET_DIR = OUTPUT_DIR / "pc_setup_guide_assets"
DOCX_PATH = OUTPUT_DIR / "Outsource_Attendance_PC_Setup_Guide.docx"

BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(17, 24, 39)
MUTED = RGBColor(75, 85, 99)
RED = RGBColor(185, 28, 28)
GREEN = RGBColor(22, 101, 52)
AMBER = RGBColor(120, 53, 15)


def font(size: int, bold: bool = False) -> ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, max_width: int, fnt: ImageFont.ImageFont) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = f"{current} {word}".strip()
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def rounded_rect(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], fill: str, outline: str = "#D1D5DB", width: int = 2) -> None:
    draw.rounded_rectangle(box, radius=14, fill=fill, outline=outline, width=width)


def draw_badge(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, fill: str = "#DC2626") -> None:
    x, y = xy
    fnt = font(28, True)
    bbox = draw.textbbox((0, 0), text, font=fnt)
    w = bbox[2] - bbox[0] + 26
    h = bbox[3] - bbox[1] + 18
    draw.ellipse((x, y, x + h, y + h), fill=fill)
    draw.text((x + 12, y + 7), text, fill="white", font=fnt)


def draw_callout(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], title: str, body: str, color: str) -> None:
    x1, y1, x2, y2 = box
    rounded_rect(draw, box, "#FFF7ED", color, 3)
    draw.text((x1 + 20, y1 + 16), title, fill=color, font=font(30, True))
    y = y1 + 56
    for line in wrap_text(draw, body, x2 - x1 - 40, font(24)):
        draw.text((x1 + 20, y), line, fill="#111827", font=font(24))
        y += 32


def save_folder_screenshot(path: Path) -> None:
    img = Image.new("RGB", (1400, 900), "#F3F4F6")
    d = ImageDraw.Draw(img)
    rounded_rect(d, (70, 70, 1330, 830), "#FFFFFF")
    d.rectangle((70, 70, 1330, 130), fill="#E5E7EB", outline="#D1D5DB")
    d.text((105, 88), "File Explorer", fill="#111827", font=font(28, True))
    rounded_rect(d, (105, 155, 1295, 205), "#F9FAFB")
    d.text((130, 168), "Desktop > OutsourceAttendance", fill="#374151", font=font(24))
    d.text((105, 235), "Name", fill="#111827", font=font(24, True))
    d.text((620, 235), "Type", fill="#111827", font=font(24, True))
    d.line((100, 270, 1295, 270), fill="#D1D5DB", width=2)
    rows = [
        ("_internal", "File folder"),
        ("OutsourceAttendance.exe", "Application"),
    ]
    y = 300
    for index, (name, typ) in enumerate(rows):
        if index == 1:
            d.rounded_rectangle((95, y - 8, 1285, y + 48), radius=8, fill="#DBEAFE", outline="#2563EB", width=3)
        icon_color = "#F59E0B" if typ == "File folder" else "#16A34A"
        d.rectangle((115, y, 155, y + 34), fill=icon_color, outline="#111827")
        d.text((175, y + 2), name, fill="#111827", font=font(25, True if index == 1 else False))
        d.text((620, y + 2), typ, fill="#374151", font=font(25))
        y += 76
    draw_badge(d, (840, 410), "1")
    draw_callout(
        d,
        (885, 385, 1265, 545),
        "Double-click here",
        "Open this file after setup is finished.",
        "#DC2626",
    )
    d.text((105, 735), "Important: copy the whole folder. Do not copy only the EXE file.", fill="#991B1B", font=font(28, True))
    img.save(path)


def save_start_screenshot(path: Path) -> None:
    img = Image.new("RGB", (1400, 900), "#E5E7EB")
    d = ImageDraw.Draw(img)
    d.rectangle((0, 780, 1400, 900), fill="#111827")
    d.rounded_rectangle((48, 806, 160, 870), radius=10, fill="#2563EB")
    d.text((76, 822), "Start", fill="white", font=font(23, True))
    rounded_rect(d, (210, 110, 1190, 735), "#FFFFFF")
    d.text((250, 150), "Windows Search", fill="#111827", font=font(34, True))
    rounded_rect(d, (250, 215, 1150, 280), "#F9FAFB", "#9CA3AF")
    d.text((280, 232), "powershell", fill="#111827", font=font(28))
    d.text((250, 325), "Best match", fill="#6B7280", font=font(24, True))
    d.rounded_rectangle((250, 365, 1150, 455), radius=12, fill="#DBEAFE", outline="#2563EB", width=3)
    d.rectangle((280, 390, 330, 430), fill="#1D4ED8")
    d.text((350, 382), "Windows PowerShell", fill="#111827", font=font(30, True))
    d.text((350, 420), "App", fill="#4B5563", font=font(22))
    draw_badge(d, (955, 352), "2")
    draw_callout(
        d,
        (985, 350, 1320, 525),
        "Click this",
        "Do not worry. It is only used one time for setup.",
        "#DC2626",
    )
    d.text((250, 610), "Shortcut: press the Windows key, type powershell, then press Enter.", fill="#111827", font=font(28, True))
    img.save(path)


def save_powershell_screenshot(path: Path) -> None:
    img = Image.new("RGB", (1400, 900), "#F3F4F6")
    d = ImageDraw.Draw(img)
    rounded_rect(d, (70, 70, 1330, 830), "#0B1020", "#1F2937", 2)
    d.rectangle((70, 70, 1330, 125), fill="#111827")
    d.text((95, 88), "Windows PowerShell", fill="#F9FAFB", font=font(26, True))
    body_font = font(23)
    green = "#86EFAC"
    text_y = 165
    commands = [
        '[Environment]::SetEnvironmentVariable("MONGODB_URI", "PASTE_MONGODB_URI_HERE", "User")',
        '[Environment]::SetEnvironmentVariable("MONGODB_DATABASE", "attendance_db", "User")',
        '[Environment]::SetEnvironmentVariable("DATALENS_ATTENDANCE_ADMIN_PASSWORD", "PASTE_ADMIN_PASSWORD_HERE", "User")',
    ]
    d.text((100, text_y), "PS C:\\Users\\User>", fill="#93C5FD", font=body_font)
    text_y += 45
    for command in commands:
        for line in wrap_text(d, command, 1160, body_font):
            d.text((100, text_y), line, fill=green, font=body_font)
            text_y += 35
        text_y += 22
    d.text((100, text_y + 15), "If there is no red error text, setup is saved.", fill="#FBBF24", font=font(26, True))
    draw_badge(d, (1080, 235), "3")
    draw_callout(
        d,
        (850, 515, 1285, 735),
        "Paste carefully",
        "Ask the admin for the real MongoDB URI and password. Do not type the placeholder text.",
        "#F97316",
    )
    img.save(path)


def save_browser_screenshot(path: Path) -> None:
    img = Image.new("RGB", (1400, 900), "#F3F4F6")
    d = ImageDraw.Draw(img)
    rounded_rect(d, (60, 55, 1340, 840), "#FFFFFF")
    d.rectangle((60, 55, 1340, 120), fill="#E5E7EB")
    rounded_rect(d, (150, 75, 1220, 105), "#FFFFFF", "#D1D5DB", 1)
    d.text((170, 78), "http://127.0.0.1:8501", fill="#374151", font=font(20))
    d.rectangle((60, 120, 340, 840), fill="#F8FAFC", outline="#D1D5DB")
    d.text((95, 155), "Attendance", fill="#111827", font=font(34, True))
    for i, item in enumerate(["Outsource Login", "Observer Desk", "Admin Panel"]):
        y = 230 + i * 90
        fill = "#FEE2E2" if i == 0 else "#FFFFFF"
        d.rounded_rectangle((90, y, 305, y + 58), radius=8, fill=fill, outline="#CBD5E1")
        d.text((110, y + 17), item, fill="#111827", font=font(20, True))
    d.text((390, 155), "Outsource Login", fill="#111827", font=font(46, True))
    d.line((390, 220, 760, 220), fill="#EF4444", width=6)
    d.rounded_rectangle((395, 275, 680, 350), radius=8, fill="#FFFFFF", outline="#CBD5E1")
    d.text((420, 292), "Current IST Shift", fill="#6B7280", font=font(18))
    d.text((420, 315), "G - General", fill="#111827", font=font(28, True))
    d.rounded_rectangle((395, 400, 1000, 620), radius=8, fill="#FFFFFF", outline="#CBD5E1")
    d.text((430, 435), "Name", fill="#111827", font=font(23, True))
    d.rounded_rectangle((430, 470, 950, 525), radius=8, fill="#F9FAFB", outline="#CBD5E1")
    d.text((450, 485), "Selected user name", fill="#6B7280", font=font(22))
    d.text((430, 545), "PC Name", fill="#111827", font=font(23, True))
    d.rounded_rectangle((540, 540, 950, 592), radius=8, fill="#F9FAFB", outline="#CBD5E1")
    d.rounded_rectangle((395, 655, 1000, 715), radius=8, fill="#EF4444", outline="#B91C1C")
    d.text((615, 672), "Submit Login", fill="#111827", font=font(24, True))
    draw_badge(d, (1040, 158), "4")
    draw_callout(
        d,
        (955, 225, 1300, 405),
        "Success",
        "When this page opens, the app is running on this PC.",
        "#16A34A",
    )
    img.save(path)


def save_troubleshooting_screenshot(path: Path) -> None:
    img = Image.new("RGB", (1400, 900), "#F3F4F6")
    d = ImageDraw.Draw(img)
    rounded_rect(d, (70, 70, 1330, 830), "#FFFFFF")
    d.text((105, 105), "If the app does not open", fill="#111827", font=font(38, True))
    d.text((105, 165), "Check these simple points before calling technical support.", fill="#4B5563", font=font(25))
    items = [
        ("Internet", "The PC must have internet because the database is online."),
        ("Whole folder", "The _internal folder must stay beside the EXE."),
        ("Restart once", "After first-time setup commands, restart the computer once."),
        ("Logs", "Open %LOCALAPPDATA%\\OutsourceAttendance\\logs\\launcher.log."),
        ("Manual URL", "If browser does not open, type http://127.0.0.1:8501."),
    ]
    y = 240
    for index, (title, body) in enumerate(items, start=1):
        d.rounded_rectangle((110, y, 1290, y + 82), radius=10, fill="#F9FAFB", outline="#D1D5DB")
        d.ellipse((135, y + 22, 175, y + 62), fill="#2563EB")
        d.text((148, y + 28), str(index), fill="white", font=font(20, True))
        d.text((200, y + 15), title, fill="#111827", font=font(25, True))
        d.text((200, y + 45), body, fill="#374151", font=font(22))
        y += 105
    img.save(path)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(table, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Outsource Attendance App")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = DARK
    r.font.name = "Calibri"
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(14)
    r2 = p2.add_run("Simple Windows PC setup guide")
    r2.font.size = Pt(14)
    r2.font.color.rgb = MUTED


def add_callout(doc: Document, title: str, body: str, fill: str, color: RGBColor) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.35)
    set_cell_margins(table, 120, 180, 120, 180)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = color
    r.font.size = Pt(11)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.add_run(body)
    doc.add_paragraph()


def add_step_heading(doc: Document, number: int, title: str) -> None:
    p = doc.add_paragraph()
    p.style = "Heading 1"
    r = p.add_run(f"Step {number}: {title}")
    r.bold = True


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.28)
        p.paragraph_format.space_after = Pt(3)
        p.add_run(f"{index}.    {item}")


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.35)
    set_cell_margins(table, 120, 160, 120, 160)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "111827")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(lines):
        if index:
            p.add_run().add_break()
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(229, 231, 235)
    doc.add_paragraph()


def add_image(doc: Document, path: Path, caption: str) -> None:
    doc.add_picture(str(path), width=Inches(6.35))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED


def build_assets() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    assets = {
        "folder": ASSET_DIR / "01_folder_exe.png",
        "start": ASSET_DIR / "02_open_powershell.png",
        "powershell": ASSET_DIR / "03_powershell_commands.png",
        "browser": ASSET_DIR / "04_browser_app.png",
        "troubleshooting": ASSET_DIR / "05_troubleshooting.png",
    }
    save_folder_screenshot(assets["folder"])
    save_start_screenshot(assets["start"])
    save_powershell_screenshot(assets["powershell"])
    save_browser_screenshot(assets["browser"])
    save_troubleshooting_screenshot(assets["troubleshooting"])
    return assets


def build_docx() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    assets = build_assets()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.1
    for style_name, size in [("Heading 1", 15), ("Heading 2", 12)]:
        styles[style_name].font.name = "Calibri"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.color.rgb = BLUE
        styles[style_name].paragraph_format.space_before = Pt(12)
        styles[style_name].paragraph_format.space_after = Pt(6)

    add_title(doc)
    add_callout(
        doc,
        "Read this first",
        "This guide is for setting up the Windows EXE version of the attendance app. The app uses the online MongoDB Atlas database, so internet is required.",
        "E8F1FF",
        BLUE,
    )

    doc.add_heading("What you need before starting", level=1)
    add_bullets(
        doc,
        [
            "The complete folder named OutsourceAttendance. Do not copy only the EXE file.",
            "The MongoDB URI given by the admin or technical person.",
            "The database name, usually attendance_db.",
            "The admin password selected by the owner.",
            "Internet connection on the PC.",
        ],
    )
    add_callout(
        doc,
        "Password safety",
        "This guide does not print the real MongoDB password. Ask the admin for the real values when setting up each PC.",
        "FFF7ED",
        AMBER,
    )

    add_step_heading(doc, 1, "Copy the app folder to the PC")
    add_numbered(
        doc,
        [
            "Copy the full OutsourceAttendance folder from USB, Google Drive, or the office shared folder.",
            "Paste the folder on the Desktop.",
            "Open the folder and check that OutsourceAttendance.exe and the internal support folder are both there.",
        ],
    )
    add_image(doc, assets["folder"], "Screenshot: the EXE and _internal folder must stay together.")

    add_step_heading(doc, 2, "Open PowerShell")
    add_numbered(
        doc,
        [
            "Click the Windows Start button at the bottom-left of the screen.",
            "Type powershell.",
            "Click Windows PowerShell.",
            "A blue or black command window will open.",
        ],
    )
    add_image(doc, assets["start"], "Screenshot: search for PowerShell from the Windows Start button.")
    add_callout(
        doc,
        "If the person cannot find Start",
        "Press the Windows key on the keyboard, type powershell, and press Enter.",
        "F0FDF4",
        GREEN,
    )

    add_step_heading(doc, 3, "Paste the setup commands")
    doc.add_paragraph("Paste one command at a time. Press Enter after each command.")
    add_code_block(
        doc,
        [
            '[Environment]::SetEnvironmentVariable("MONGODB_URI", "PASTE_MONGODB_URI_HERE", "User")',
            '[Environment]::SetEnvironmentVariable("MONGODB_DATABASE", "attendance_db", "User")',
            '[Environment]::SetEnvironmentVariable("DATALENS_ATTENDANCE_ADMIN_PASSWORD", "PASTE_ADMIN_PASSWORD_HERE", "User")',
        ],
    )
    add_bullets(
        doc,
        [
            "Replace PASTE_MONGODB_URI_HERE with the real MongoDB URI.",
            "Replace PASTE_ADMIN_PASSWORD_HERE with the real admin password.",
            "If no red error message appears, the command is saved.",
            "After all three commands are done, close PowerShell and restart the computer once.",
        ],
    )
    add_image(doc, assets["powershell"], "Screenshot: PowerShell window with the three setup commands.")

    add_step_heading(doc, 4, "Start the attendance app")
    add_numbered(
        doc,
        [
            "After restarting, open the OutsourceAttendance folder on the Desktop.",
            "Double-click OutsourceAttendance.exe.",
            "Wait 10 to 30 seconds. The browser should open automatically.",
            "If the browser does not open, open Chrome or Edge and type http://127.0.0.1:8501.",
        ],
    )
    add_image(doc, assets["browser"], "Screenshot: successful app page in the browser.")

    add_step_heading(doc, 5, "Daily use")
    add_bullets(
        doc,
        [
            "For daily use, only double-click OutsourceAttendance.exe.",
            "Do not delete the internal support folder.",
            "Do not move the EXE alone to another folder.",
            "If Windows shows SmartScreen, click More info, then Run anyway only if the file came from the office admin.",
            "To stop the app, close the browser tab. If needed, restart the PC.",
        ],
    )

    add_step_heading(doc, 6, "Troubleshooting")
    add_image(doc, assets["troubleshooting"], "Screenshot: common things to check if the app does not open.")
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.35)
    set_cell_margins(table, 90, 120, 90, 120)
    hdr = table.rows[0].cells
    hdr[0].text = "Problem"
    hdr[1].text = "What to do"
    for cell in hdr:
        set_cell_shading(cell, "F2F4F7")
        for paragraph in cell.paragraphs:
            paragraph.runs[0].bold = True
    rows = [
        ("App does not open", "Double-click again and wait 30 seconds. Then try http://127.0.0.1:8501 in Chrome or Edge."),
        ("MongoDB error", "Check internet. Then ask admin to confirm MongoDB URI and Atlas Network Access."),
        ("Password not working", "Ask admin to confirm DATALENS_ATTENDANCE_ADMIN_PASSWORD."),
        ("EXE moved alone", "Copy the full OutsourceAttendance folder again."),
        ("Need log file", "Open %LOCALAPPDATA%\\OutsourceAttendance\\logs\\launcher.log."),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()

    add_callout(
        doc,
        "Important for admin",
        "All PCs use the same online database. If many PCs are used, MongoDB Atlas Network Access must allow that office internet connection.",
        "FEE2E2",
        RED,
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Outsource Attendance App setup guide")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
    print(DOCX_PATH)
